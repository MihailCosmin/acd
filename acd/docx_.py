import sys
from platform import system


from os import walk
from os import mkdir
from os import remove
from os import rename
from os import listdir
if system() == "Windows":
    from os import startfile
from os.path import sep
from os.path import join
from os.path import isdir
from os.path import exists
from os.path import isfile
from os.path import dirname
from os.path import basename

from subprocess import Popen

from re import findall
from re import search

from shutil import copy
from shutil import rmtree

from traceback import format_exc

from docx import Document

from time import sleep
from time import time

from lxml import etree

from traceback import format_exc
if system() == "Windows":
    from win32com.client import Dispatch  # Default .docx to .pdf solution

from tqdm import tqdm

# import aspose.words as aw  # Proprietary library - needs license > 1000 USD (currently adds watermark and truncates some pages)
# from docx2pdf import convert  # This library uses win32com

import sys

if sys.version_info >= (3, 12):
    from PySide6.QtWidgets import QMainWindow
    from PySide6.QtCore import Signal
else:
    from PySide2.QtWidgets import QMainWindow
    from PySide2.QtCore import Signal

from typing import List
from docx import Document

from .archive import unarchive_file
from .archive import zip_word_folder


WORD_EXTENSIONS = [".docx", ".docm", ".doc", ".dotx", ".dotm", ".dot", ".docb"]
WORD_NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
}

def read_word_footers(docx: str, back_up: bool = False, debug: bool = False) -> str:
    """Read all footers in a docx file or docm file

    Args:
        docx (str): path + filename of the docx file or docm file
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
    """
    output = ""

    unarchive_file(docx)
    ext = "." + docx.split(".")[-1]
    docx_dir = docx.replace(ext, "")

    try:
        footers = 0
        for xml in listdir(join(docx_dir, "word")):
            if xml.startswith("footer"):
                footers += 1
        for footer in range(1, footers + 1, 1):
            xml_path = join(docx_dir, "word", f"footer{footer}.xml")
            xml_tree = etree.parse(xml_path)
            for ind, paragraph in enumerate(xml_tree.xpath("w:p", namespaces=WORD_NS), start=1):
                try:
                    para_value = ""
                    for w_r in paragraph.xpath("w:r", namespaces=WORD_NS):
                        for w_t in w_r.xpath("w:t", namespaces=WORD_NS):
                            try:
                                para_value += w_t.text
                            except IndexError:
                                para_value += ""
                    # print(f"Footer {footer}, paragraph {ind}: {para_value}")
                    output += f"Footer {footer}, paragraph {ind}: {para_value}\n"
                    if "Liebherr-Elektronik GmbH" in para_value:
                        pass
                except IndexError:
                    # print(f"Footer {footer}, paragraph {ind}: No text")
                    output += f"Footer {footer}, paragraph {ind}: No text\n"
            for table_no, table in enumerate(xml_tree.xpath("w:tbl", namespaces=WORD_NS), start=1):
                for row_no, row in enumerate(table.xpath("w:tr", namespaces=WORD_NS)):
                    for cell_no, cell in enumerate(row.xpath("w:tc", namespaces=WORD_NS), start=1):
                        cell_value = ""
                        for paragraph in cell.xpath("w:p", namespaces=WORD_NS):
                            for w_r in paragraph.xpath("w:r", namespaces=WORD_NS):
                                for w_t in w_r.xpath("w:t", namespaces=WORD_NS):
                                    try:
                                        cell_value += w_t.text
                                    except IndexError:
                                        cell_value += ""
                        # print(f"Footer: {footer} Tabel {table_no}, row {row_no}, cell {cell_no}: {cell_value}")
                        output += f"Footer: {footer} Tabel {table_no}, row {row_no}, cell {cell_no}: {cell_value}\n"
    except Exception as err:
        # print(f"Error, could not read footers: {err}\n{format_exc()}")
        output = f"Error, could not read footers: {err}\n{format_exc()}"
        return output
    return output



def get_regex_string(docx: str, reg_ex: str, back_up: bool = False, debug: bool = False) -> int:
    """Read all footers in a docx file or docm file

    Args:
        docx (str): path + filename of the docx file or docm file
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
    """

    unarchive_file(docx)
    ext = "." + docx.split(".")[-1]
    docx_dir = docx.replace(ext, "")

    try:
        footers = 0
        for xml in listdir(join(docx_dir, "word")):
            if xml.startswith("footer"):
                footers += 1
        for footer in range(1, footers + 1, 1):
            xml_path = join(docx_dir, "word", f"footer{footer}.xml")
            xml_tree = etree.parse(xml_path)
            for ind, paragraph in enumerate(xml_tree.xpath(".//w:p", namespaces=WORD_NS), start=1):
                try:
                    para_value = ""
                    for w_r in paragraph.xpath(".//w:r", namespaces=WORD_NS):
                        for w_t in w_r.xpath(".//w:t", namespaces=WORD_NS):
                            try:
                                para_value += w_t.text
                            except IndexError:
                                para_value += ""
                    if debug:
                        print(f"Footer {footer}, paragraph {ind}: {para_value}")
                    if search(reg_ex, para_value):
                        return search(reg_ex, para_value).group(0)
                except IndexError:
                    print(f"Footer {footer}, paragraph {ind}: No text")
            for table_no, table in enumerate(xml_tree.xpath(".//w:tbl", namespaces=WORD_NS), start=1):
                for row_no, row in enumerate(table.xpath(".//w:tr", namespaces=WORD_NS)):
                    for cell_no, cell in enumerate(row.xpath(".//w:tc", namespaces=WORD_NS), start=1):
                        cell_value = ""
                        for paragraph in cell.xpath(".//w:p", namespaces=WORD_NS):
                            for w_r in paragraph.xpath(".//w:r", namespaces=WORD_NS):
                                for w_t in w_r.xpath(".//w:t", namespaces=WORD_NS):
                                    try:
                                        cell_value += w_t.text
                                    except IndexError:
                                        cell_value += ""
                        if debug:
                            print(f"Footer: {footer} Tabel {table_no}, row {row_no}, cell {cell_no}: {cell_value}")
                        # cell_value = cell_value.replace("Copyright", "©").replace("copyright", "©")
                        if search(reg_ex, cell_value):
                            return search(reg_ex, cell_value).group(0)
    except Exception as err:
        print(f"Error, could not read footers: {err}\n{format_exc()}")
        return 1
    return 0

from typing import List
from docx import Document


def replace_copyright(docx: str, debug: bool = False) -> int:
    """Replace the string "Copyright" with the character '©'

    Args:
        docx (str): path + filename of the docx file or docm file
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
    """

    unarchive_file(docx)
    ext = "." + docx.split(".")[-1]
    docx_dir = docx.replace(ext, "")

    try:
        footers = 0
        for xml in listdir(join(docx_dir, "word")):
            if xml.startswith("footer"):
                footers += 1

        for footer in range(1, footers + 1):
            xml_path = join(docx_dir, "word", f"footer{footer}.xml")
            xml_tree = etree.parse(xml_path)
            
            for ind, paragraph in enumerate(xml_tree.xpath(".//w:p", namespaces=WORD_NS), start=1):
                try:
                    para_value = ""
                    for w_r in paragraph.xpath(".//w:r", namespaces=WORD_NS):
                        for w_t in w_r.xpath(".//w:t", namespaces=WORD_NS):
                            try:
                                para_value += w_t.text
                            except IndexError:
                                para_value += ""

                    # Replace "Copyright" with '©' in the paragraph
                    para_value = para_value.replace("Copyright", "©")
                    # Clear the existing text in the paragraph

                except IndexError:
                    print(f"Footer {footer}, paragraph {ind}: No text")

            for table_no, table in enumerate(xml_tree.xpath(".//w:tbl", namespaces=WORD_NS), start=1):
                for row_no, row in enumerate(table.xpath(".//w:tr", namespaces=WORD_NS)):
                    for cell_no, cell in enumerate(row.xpath(".//w:tc", namespaces=WORD_NS), start=1):
                        cell_value = ""
                        for paragraph in cell.xpath(".//w:p", namespaces=WORD_NS):
                            for w_r in paragraph.xpath(".//w:r", namespaces=WORD_NS):
                                for w_t in w_r.xpath(".//w:t", namespaces=WORD_NS):
                                    try:
                                        if "Copyright" in w_t.text:
                                            w_t.text = w_t.text.replace("Copyright", "©")
                                        cell_value += w_t.text
                                    except IndexError:
                                        cell_value += ""

                        # Replace "Copyright" with '©' in the cell
                        print(cell_value)
                        cell_value = cell_value.replace("Copyright", "©")
                        print(cell_value)
        xml_tree.write(xml_path, encoding="utf-8", pretty_print=True)

    except Exception as err:
        print(f"Error, could not read footers: {err}\n{format_exc()}")
        return 1

    return 0


def get_template_version(docx: str, reg_ex: str, back_up: bool = False, debug: bool = False) -> any:
    """Read all footers in a docx file or docm file

    Args:
        docx (str): path + filename of the docx file or docm file
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
    """

    unarchive_file(docx)
    ext = "." + docx.split(".")[-1]
    docx_dir = docx.replace(ext, "")

    try:
        footers = 0
        for xml in listdir(join(docx_dir, "word")):
            if xml.startswith("footer"):
                footers += 1
        for footer in range(1, footers + 1, 1):
            xml_path = join(docx_dir, "word", f"footer{footer}.xml")
            xml_tree = etree.parse(xml_path)
            for ind, paragraph in enumerate(xml_tree.xpath(".//w:p", namespaces=WORD_NS), start=1):
                try:
                    para_value = ""
                    for w_r in paragraph.xpath(".//w:r", namespaces=WORD_NS):
                        for w_t in w_r.xpath(".//w:t", namespaces=WORD_NS):
                            try:
                                para_value += w_t.text
                            except IndexError:
                                para_value += ""
                    if debug:
                        print(f"Footer {footer}, paragraph {ind}: {para_value}")
                    if search(reg_ex, para_value):
                        return search(reg_ex, para_value).group(0)
                except IndexError:
                    print(f"Footer {footer}, paragraph {ind}: No text")
            for table_no, table in enumerate(xml_tree.xpath(".//w:tbl", namespaces=WORD_NS), start=1):
                for row_no, row in enumerate(table.xpath(".//w:tr", namespaces=WORD_NS)):
                    for cell_no, cell in enumerate(row.xpath(".//w:tc", namespaces=WORD_NS), start=1):
                        cell_value = ""
                        for paragraph in cell.xpath(".//w:p", namespaces=WORD_NS):
                            for w_r in paragraph.xpath(".//w:r", namespaces=WORD_NS):
                                for w_t in w_r.xpath(".//w:t", namespaces=WORD_NS):
                                    try:
                                        cell_value += w_t.text
                                    except IndexError:
                                        cell_value += ""
                        if debug:
                            print(f"Footer: {footer} Tabel {table_no}, row {row_no}, cell {cell_no}: {cell_value}")
                        cell_value = cell_value.replace("Copyright", "©").replace("copyright", "©")
                        if cell_no == 2 and cell_value != "":
                            return cell_value
    except Exception as err:
        print(f"Error, could not read footers: {err}\n{format_exc()}")
        return 1
    return 0



def docx_footer_replace(docx: str, old: str, new: str, back_up: bool = False, debug: bool = False) -> int:
    """Replace a string in the footer of a docx file or docm file

    Args:
        docx (str): path + filename of the docx file or docm file
        old (str): string to replace
        new (str): new string
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
    
    NOTE: Carefull this breaks word variables!!
    """
    if ".docx" in docx.lower():
        doc_type = ".docx"
    elif ".docm" in docx.lower():  # Cosmin changed. Was elif ".docm" not in docx.lower():
        doc_type = ".docm"
    else:
        print(f"Error, file is not a docx or docm file: {docx}")
        return 1

    try:
        doc = Document(docx)
        for secn_num, section in enumerate(doc.sections, start=1):
            for footer_para in section.footer.paragraphs:
                if debug:
                    print(f"Section {secn_num}, Footer para text: {footer_para.text}")
                if old in footer_para.text:
                    footer_para.text = footer_para.text.replace(old, new)
            for footer_table in section.footer.tables:
                for row in footer_table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if debug:
                                print(f"Section {secn_num}, Footer table cell text: {para.text}")
                            if old in para.text:
                                para.text = para.text.replace(old, new)
    except Exception as err:
        print(f"Error: {err}\n{format_exc()}")
        return 1
    try:
        if back_up:
            copy(docx, docx.replace(doc_type, f"_backup{doc_type}"))
    except Exception as err:
        print(f"Error, could not make back-up: {err}\n{format_exc()}")
        return 1

    try:
        doc.save(docx)
    except Exception as err:
        print(f"Error, could not save document: {err}\n{format_exc()}")
        return 1
    return 0


def docx_header_replace(docx: str, old: str, new: str, back_up: bool = False, debug: bool = False) -> int:
    """Replace a string in the header of a docx file or docm file

    Args:
        docx (str): path + filename of the docx file or docm file
        old (str): string to replace
        new (str): new string
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
        
    NOTE: Carefull this breaks word variables!!
    """
    if ".docx" in docx.lower():
        doc_type = ".docx"
    elif ".docm" in docx.lower():  # Cosmin changed. Was elif ".docm" not in docx.lower():
        doc_type = ".docm"
    else:
        print(f"Error, file is not a docx or docm file: {docx}")
        return 1

    try:
        doc = Document(docx)

        # If something is not printed here, it might not be a header
        for secn_num, section in enumerate(doc.sections, start=1):
            for header_para in section.header.paragraphs:
                if debug:
                    print(f"Section {secn_num}, Header para text: {header_para.text}")
                if old in header_para.text:
                    header_para.text = header_para.text.replace(old, new)

            for header_table in section.header.tables:
                for row in header_table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if debug:
                                print(f"Section {secn_num}, Header table cell text: {para.text}")
                            if old in para.text:
                                para.text = para.text.replace(old, new)
    except Exception as err:
        print(f"Error: {err}\n{format_exc()}")
        return 1
    try:
        if back_up:
            copy(docx, docx.replace(doc_type, f"_backup{doc_type}"))
    except Exception as err:
        print(f"Error, could not make back-up: {err}\n{format_exc()}")
        return 1

    try:
        doc.save(docx)
    except Exception as err:
        print(f"Error, could not save document: {err}\n{format_exc()}")
        return 1
    return 0


def docx_content_replace(docx: str, old: str, new: str, back_up: bool = False, debug: bool = False) -> int:
    """Replace a string in the content of a docx file or docm file

    Args:
        docx (str): path + filename of the docx file or docm file 
        old (str): string to replace
        new (str): new string
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
    
    NOTE: Carefull this breaks word variables!!
    """
    if ".docx" in docx.lower():
        doc_type = ".docx"
    elif ".docm" in docx.lower():  # Cosmin changed. Was elif ".docm" not in docx.lower():
        doc_type = ".docm"
    else:
        print(f"Error, file is not a docx or docm file: {docx}")
        return 1

    try:
        doc = Document(docx)
        for para in doc.paragraphs:
            if debug:
                print(f"Para text: {para.text}")
            if old in para.text:
                para.text = para.text.replace(old, new)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if debug:
                            print(f"Table cell text: {para.text}")
                        if old in para.text:
                            para.text = para.text.replace(old, new)
    except Exception as err:
        print(f"Error: {err}\n{format_exc()}")
        return 1
    try:
        if back_up:
            copy(docx, docx.replace(doc_type, f"_backup{doc_type}"))
    except Exception as err:
        print(f"Error, could not make back-up: {err}\n{format_exc()}")
        return 1

    try:
        doc.save(docx)
    except Exception as err:
        print(f"Error, could not save document: {err}\n{format_exc()}")
        return 1
    return 0

def replace_media(docx: str, old_media_name: str, new_media_path: str, back_up: bool = False, debug: bool = False) -> int:
    """Replace a media file in a docx file or docm file

    Args:
        docx (str): path + filename of the docx file or docm file
        old_media_name (str): name of the media file to replace
        new_media_path (str): path + filename of the new media file
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not
    """
    if ".docx" in docx.lower():
        doc_type = ".docx"
    elif ".docm" in docx.lower():  # Cosmin changed. Was elif ".docm" not in docx.lower():
        doc_type = ".docm"
    else:
        print(f"Error, file is not a docx or docm file: {docx}")
        return 1

    unarchive_file(docx)
    docx_dir = docx.replace(doc_type, "")
    try:
        if isfile(join(docx_dir, "word", "media", old_media_name)):
            remove(join(docx_dir, "word", "media", old_media_name))
            copy(new_media_path, join(docx_dir, "word", "media", old_media_name))
    except Exception as err:
        print(f"Error, could not replace media: {err}\n{format_exc()}")
        return 1
    try:
        if back_up:
            copy(docx, docx.replace(doc_type, f"_backup{doc_type}"))
    except Exception as err:
        print(f"Error, could not make back-up: {err}\n{format_exc()}")
        return 1
    try:
        zip_word_folder(docx_dir)
        remove(docx)
        rename(join(docx.replace(doc_type, ".zip")), docx)
    except Exception as err:
        print(f"Error, could not zip folder: {err}\n{format_exc()}")
        return 1
    return 0

def word2pdf(
        dir_name: str,
        bookmarks: int = 0,
        skip_existing: bool = False,
        debug: bool = False,
        qt_window: QMainWindow = None,
        progress: Signal = Signal(0),
        console: Signal = Signal("")):
    """Convert word files to PDF

    Args:
        dir_name (str): folder path
        bookmarks (int): bookmarks or not
        skip_existing (bool): skip existing files or not
        debug (bool, optional): Print debug info or not. Defaults to False.
        qt_window (QMainWindow, optional): QT Main window. Defaults to None.
    """
    errors = 0

    list_of_files = []

    for (dirpath, _, filenames) in walk(dir_name):
        list_of_files += [join(dirpath, file_).replace("\\", sep).replace("/", sep) for file_ in filenames]

    progress_index = 0
    progress_max = len(list_of_files)
    # Cosmin 14.02.2022 - Created below container because if we print a progressbar with tqdm
    # when the console is hidden for exe, the script doesn't run.
    in_container = list_of_files if qt_window is not None else tqdm(
        list_of_files)

    start_time = time()

    for file_ in in_container:
        ext = "." + file_.split(".")[-1].lower()
        if qt_window is not None:
            progress.emit(progress_index * 100 / progress_max)
            progress_index += 1
        try:
            if ext in WORD_EXTENSIONS and "$" not in file_:
                if debug:
                    console.emit(f"Converting {file_.split(sep)[-1]}")
                pdf_file = file_.replace(ext, ".pdf")

                if isfile(pdf_file) and not skip_existing:
                    remove(pdf_file)
                elif isfile(pdf_file) and skip_existing:
                    continue
                word = Dispatch("Word.Application")

                doc = word.Documents.Open(file_)

                # Save as PDF - Advanced Option, with choices
                doc.ExportAsFixedFormat(pdf_file,
                                        17,  # ExportFormat # 17 - wdExportFormatPDF, 18 - wdExportFormatXPS
                                        False,  # OpenAfterExport # Boolean value
                                        0,  # OptimizeFor # 0 - wdExportOptimizeForPrint, 1 - wdExportOptimizeForOnScreen
                                        0,  # Range # 0 - wdExportAllDocument,
                                            # 1 - wdExportSelection,
                                            # 2 - wdExportCurrentPage,
                                            # 3 - wdExportFromTo
                                        1,  # From # keep 1 as default if Range is 0
                                        1,  # To # keep 1 as default if Range is 0
                                        7,  # Item # 0 - wdExportDocumentContent, 7 - wdExportDocumentWithMarkup
                                        False,  # IncludeDocProps # Boolean value
                                        False,  # KeepIRM # Boolean value
                                        bookmarks,  # CreateBookmarks
                                                    # 0 - wdExportCreateNoBookmarks,
                                                    # 1 - wdExportCreateHeadingBookmarks,
                                                    # 2 - wdExportCreateWordBookmarks
                                        True,  # DocStructureTags # Boolean value
                                        True,  # BitmapMissingFonts # Boolean value
                                        False,  # UseISO19005_1 # Boolean value
                                        )

                doc.Close()

        except (OSError, RuntimeError, ValueError, NameError, AttributeError):
            if debug:
                if not isdir(join(qt_window.exe_path, "debug")):
                    mkdir(join(qt_window.exe_path, "debug"))
                errors += 1
                if qt_window is not None:
                    console.emit(f"There was an error while processing this file: {file_}")
                    console.emit("The file was not saved to PDF. Please correct the issue and try again.")
                    console.emit(f"The error message was: {format_exc()}")

                if exists(join(qt_window.exe_path, "debug", "word2pdf_log.txt")):
                    with open(join(qt_window.exe_path, "debug", "word2pdf_log.txt"), "a", encoding='utf-8') as log:
                        log.write(f"There was an error while processing this file: {file_}\n")
                        log.write("The file was not saved to PDF. Please correct the issue and try again.\n")
                        log.write(f"The error message was: {format_exc()}\n\n")
                else:
                    with open(join(qt_window.exe_path, "debug", "word2pdf_log.txt"), "w", encoding='utf-8') as log:
                        log.write(f"There was an error while processing this file: {file_}\n")
                        log.write("The file was not saved to PDF. Please correct the issue and try again.\n")
                        log.write(f"The error message was: {format_exc()}\n\n")
                # word.Quit()
                doc.Close()
                continue

    if qt_window is None:
        if errors == 0:
            print("Process finished with no errors!")
            sleep(5)
        else:
            PLURAL = "" if errors == 1 else "s"
            print(f"Process finished with {errors} error{PLURAL}!\n")
            while True:
                ANSWER = str(input('Would you like to open the error log? (y/n)\n'))
                if ANSWER in ('y', 'n', 'Y', 'N'):
                    break
                print("Invalid input. \n")
            if ANSWER == 'y' or ANSWER == 'Y':
                if system() == 'Windows':
                    startfile(join(qt_window.exe_path, "debug", "word2pdf_log.txt"))
                elif system() == 'Linux':
                    Popen(['xdg-open', join(qt_window.exe_path, "debug", "word2pdf_log.txt")])
    elif qt_window is not None:
        progress.emit(100)
        console.emit(f"Finished in {round(time() - start_time, 2)} seconds")
        if errors > 0:
            return 1
        return 0

def get_table_column_widths(docx: str, where: str) -> dict:
    """Get all tables column widths in a docx file
    You need to mention where to search for the tables: header, footer or content

    Args:
        docx (str): path + filename of the docx file or docm file
        where (str): where to search for the tables: header, footer or content

    Returns:
        dict: Dictionary with the table names as keys and the column widths as values
    """
    doc = Document(docx)
    widths = {}
    if where == "content":
        for table_no, table in enumerate(doc.tables):
            current_table_widths = []
            for row in table.rows:
                for cell in row.cells:
                    current_table_widths.append(cell.width)
            if len(current_table_widths) > 0 and "content_table" not in widths:
                widths[f"content_table_{table_no}"] = current_table_widths
    else:
        for sec_num, section in enumerate(doc.sections, start=1):
            container = section.footer.tables if where == "footer" else section.header.tables if where == "header" else None
            for table in container:
                current_table_widths = []
                for row in table.rows:
                    for cell in row.cells:
                        current_table_widths.append(cell.width)
                if len(current_table_widths) > 0 and f"section_{sec_num}_{where}_table" not in widths:
                    widths[f"section_{sec_num}_{where}_table"] = current_table_widths
    return widths

def adjust_column_widths(
        docx: str, where: str, widths: list = None, col: int = None,
        difference: float = None, xml_method: bool = True,
        back_up: bool = False, debug: bool = False) -> int:
    """Adjust the column widths of all tables in a docx file
    You need to mention where to search for the tables: header, footer or content

    Args:
        docx (str): path + filename of the docx file or docm file
        where (str): where to search for the tables: header, footer or content
        widths (list, optional): list of widths to set. Defaults to None.
        col (int, optional): column to adjust. Defaults to None.
        difference (float, optional): difference to add to the column width as a float of around 1. Defaults to None.
            Ex: 0.9 will reduce the size of the column by 10%
                1.1 will increase the size of the column by 10%
        xml_method (bool, optional): use XML manipulation instead of python-docx. Defaults to True.
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.

    Returns:
        int: 0 if successful, 1 if not


    NOTE: Manipulations with python-docx doesn't seem to work as expected.
    Use XML manipulation instead.
    """

    if xml_method:
        col_width = None
        w_attr = f"{{{WORD_NS['w']}}}w"
        w_left = f"{{{WORD_NS['w']}}}left"
        w_right = f"{{{WORD_NS['w']}}}right"
        unarchive_file(docx)
        ext = "." + docx.split(".")[-1]
        docx_dir = docx.replace(ext, "")
        if where in ("header", "footer"):
            for xml in listdir(join(docx_dir, "word")):
                if xml.startswith(where):
                    xml_path = join(docx_dir, "word", xml)
                    xml_tree = etree.parse(xml_path)
                    for table_xml in xml_tree.xpath("//w:tbl", namespaces=WORD_NS):
                        for table_grid in table_xml.xpath("w:tblGrid", namespaces=WORD_NS):
                            for ind, column in enumerate(table_grid.xpath("w:gridCol", namespaces=WORD_NS), start=1):
                                if int(float(column.attrib[w_attr])) >= 4800:  # Cosmin added first conversion to float, as some attrib have 4151.59999
                                    break
                                if ind == col:
                                    if difference is not None:
                                        # delta = int((int(float(column.attrib[w_attr])) - (int(float(column.attrib[w_attr])) * difference)) / 2)
                                        delta = int((int(float(column.attrib[w_attr])) - 4800) / 2)
                                        # col_width = str(int(float(column.attrib[w_attr])) * difference)
                                        col_width = "4800"
                                        column.attrib[w_attr] = col_width
                                        prev_col_width = str(int(float(table_grid.xpath(f"w:gridCol[{ind - 1}]", namespaces=WORD_NS)[0].attrib[w_attr])) + int(delta))
                                        table_grid.xpath(f"w:gridCol[{ind - 1}]", namespaces=WORD_NS)[0].attrib[w_attr] = prev_col_width
                                        next_col_width = str(int(float(table_grid.xpath(f"w:gridCol[{ind + 1}]", namespaces=WORD_NS)[0].attrib[w_attr])) + int(delta))
                                        table_grid.xpath(f"w:gridCol[{ind + 1}]", namespaces=WORD_NS)[0].attrib[w_attr] = next_col_width
                                    break
                        # if col_width is None:
                        #     break
                        for table_row in table_xml.xpath(".//w:tr", namespaces=WORD_NS):
                            for ind, row_cell in enumerate(table_row.xpath(".//w:tc", namespaces=WORD_NS), start=1):
                                if ind == col - 1:
                                    if col_width is not None:
                                        row_cell.xpath(".//w:tcPr/w:tcW", namespaces=WORD_NS)[0].attrib[w_attr] = prev_col_width
                                elif ind == col:
                                    if col_width is not None:
                                        row_cell.xpath(".//w:tcPr/w:tcW", namespaces=WORD_NS)[0].attrib[w_attr] = col_width
                                    for w_p in row_cell.xpath(".//w:p", namespaces=WORD_NS):
                                        for w_r in w_p.xpath(".//w:r", namespaces=WORD_NS):
                                            # Insert xml as child of w_r
                                            for prefix, uri in WORD_NS.items():
                                                etree.register_namespace(prefix, uri)
                                            xml_elem = '<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:b w:val="0"/><w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>'
                                            xml_elem = etree.fromstring(xml_elem)
                                            w_r.insert(0, xml_elem)
                                elif ind == col + 1:
                                    if col_width is not None:
                                        row_cell.xpath(".//w:tcPr/w:tcW", namespaces=WORD_NS)[0].attrib[w_attr] = next_col_width
                                    for para in row_cell.xpath(".//w:p", namespaces=WORD_NS):
                                        for para_prop in para.xpath(".//w:pPr", namespaces=WORD_NS):
                                            for ind_prop in para_prop.xpath(".//w:ind", namespaces=WORD_NS):
                                                ind_prop.attrib[w_left] = str(int(float(ind_prop.attrib[w_left])) + delta)
                                                ind_prop.attrib[w_right] = "0"
                    xml_tree.write(xml_path, encoding="utf-8", xml_declaration=True)
        else:
            print("Content not implemented yet")
        try:
            if back_up:
                copy(docx, docx.replace(ext, f"_backup{ext}"))
            zip_word_folder(docx_dir)
            rmtree(docx_dir)
            deleted = False
            while not deleted:
                try:
                    remove(docx)
                    deleted = True
                except Exception as err:
                    print(f"Error, could not delete file: {err}\n{format_exc()}")
                    sleep(1)
                
            rename(docx.replace(ext, ".zip"), docx)
        except Exception as err:
            print(f"Error, could not zip folder: {err}\n{format_exc()}")
            return 1
        return 0

    if not xml_method:
        doc = Document(docx)
        if where == "content":
            for table_no, table in enumerate(doc.tables):
                for row in table.rows:
                    for ind, cell in enumerate(row.cells):
                        cell.width = widths[table_no]
        else:
            for sec_num, section in enumerate(doc.sections, start=1):
                if debug:
                    print(f"Section {sec_num}")
                container = section.footer.tables if where == "footer" else section.header.tables if where == "header" else None
                for table in container:
                    for row in table.rows:
                        for ind, cell in enumerate(row.cells):
                            if "Liebherr-Elektronik GmbH" in cell.text:  # Targeted processing for specific keywords
                                print(f"Cell 1 {ind} has text: {cell.text} and width: {cell.width}")
                                cell.width = cell.width + (row.cells[ind - 1].width * 0.3) + (row.cells[ind + 1].width * 0.3)
                                row.cells[ind - 1].width = row.cells[ind - 1].width * 0.7
                                row.cells[ind + 1].width = row.cells[ind + 1].width * 0.7
                                print(f"Cell 2 {ind} has text: {cell.text} and width: {cell.width}")
                                break
                            else:
                                print(f"Cell {ind} has width: {cell.width}")

                            # cell.width = widths[ind]  # Adjust all widths
        try:
            if back_up:
                copy(docx, docx.replace(".docx", "_backup.docx"))
            doc.save(docx)
        except Exception as err:
            print(f"Error, could not save document: {err}\n{format_exc()}")
            return 1
        return 0


def get_footer_type(docx: str) -> str:
    """Get the footer type of a docx file

    Args:
        docx (str): path + filename of the docx file or docm file

    Returns:
        str: footer type
    """
    doc = Document(docx)
    for section in doc.sections:
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.width > 600000 and cell.width < 700000:
                        return "old_tpye"
                    if cell.width > 1000000:
                        return "new_type"
                    return "unknown"

def update_footer_table_widths(
        docx: str, widths: list = None,
        back_up: bool = False, debug: bool = False):
    """Update the footer table widths of a docx file

    Args:
        docx (str): path + filename of the docx file or docm file
        widths (list, optional): list of widths to set. Defaults to None.
        back_up (bool, optional): create a backup of the original file. Defaults to False.
        debug (bool, optional): print debug messages. Defaults to False.
    """
    if get_footer_type(docx) == "old_type":
        print("Yo 1")
        adjust_column_widths(
            docx,
            "footer",
            [608965, 651510, 3690620, 629920, 720090, 4951095, 4951095, 4951095, 629920, 720090] if widths is None else widths,
            back_up=back_up,
            debug=debug
        )
    elif get_footer_type(docx) == "new_type":
        print("Yo 2")
        adjust_column_widths(
            docx,
            "footer",
            [802005, 1854200, 3600450, 1753235, 970915, 8980805, 8980805, 8980805, 8980805, 8980805] if widths is None else widths,
            back_up=back_up,
            debug=debug
        )
    else:
        print("Unknown footer type")
